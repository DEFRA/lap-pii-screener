from __future__ import annotations

import asyncio
import base64
import bz2
import csv
import fnmatch
import gzip
import io
import re
import sys
import tarfile
import urllib.parse
import zipfile
from pathlib import Path
from typing import Callable, Iterator, Optional

import aiofiles

from models.finding import Finding, ScanConfig
from remediation.engine import RemediationEngine
from remediation.regulation_engine import RegulationEngine
from scanners.base import AbstractScanner

_ENGINE = RemediationEngine()
_REG_ENGINE = RegulationEngine()

# ─── Optional third-party extractors ─────────────────────────────────────────

_DOCX_AVAILABLE = False
_XLSX_AVAILABLE = False
_XLS_AVAILABLE = False
_RTF_AVAILABLE = False
_DOC_AVAILABLE = False
_PDF_AVAILABLE = False
_MSG_AVAILABLE = False
_PYARROW_AVAILABLE = False  # handles both .parquet and .orc
_AVRO_AVAILABLE = False

try:
    import docx as _docx_module  # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    import openpyxl as _openpyxl  # openpyxl
    _XLSX_AVAILABLE = True
except ImportError:
    pass

try:
    import xlrd as _xlrd  # xlrd (legacy .xls)
    _XLS_AVAILABLE = True
except ImportError:
    pass

try:
    from striprtf.striprtf import rtf_to_text as _rtf_to_text  # striprtf
    _RTF_AVAILABLE = True
except ImportError:
    pass

try:
    import olefile as _olefile  # olefile (legacy .doc)
    _DOC_AVAILABLE = True
except ImportError:
    pass

try:
    from pdfminer.high_level import extract_text as _pdfminer_extract  # pdfminer.six
    _PDF_AVAILABLE = True
except ImportError:
    pass

try:
    import extract_msg as _extract_msg_lib  # extract-msg (Outlook .msg)
    _MSG_AVAILABLE = True
except ImportError:
    pass

try:
    import pyarrow.parquet as _pq  # pyarrow (.parquet)
    import pyarrow.orc as _pa_orc  # pyarrow (.orc)
    _PYARROW_AVAILABLE = True
except ImportError:
    pass

try:
    import fastavro as _fastavro  # fastavro (.avro)
    _AVRO_AVAILABLE = True
except ImportError:
    pass

# Extensions that require binary-to-text extraction before scanning
_BINARY_DOC_EXTENSIONS: set[str] = {
    ".docx", ".xlsx", ".xls", ".rtf", ".eml",
    ".doc", ".pdf", ".msg", ".parquet", ".orc", ".avro",
}

# Archive file extensions handled by the traversal engine
_ARCHIVE_EXTENSIONS: set[str] = {".zip", ".gz", ".bz2", ".tar", ".tgz"}

# Maximum nested archive depth (FR-ING-004 default: 10)
_MAX_ARCHIVE_DEPTH: int = 10

# ─── Check-digit / format validators ────────────────────────────────────────

def _luhn(number: str) -> bool:
    """Luhn mod-10 check for payment card numbers."""
    digits = [int(d) for d in number if d.isdigit()]
    if len(digits) < 13:
        return False
    total = 0
    for i, d in enumerate(reversed(digits)):
        if i % 2 == 1:
            d *= 2
            if d > 9:
                d -= 9
        total += d
    return total % 10 == 0


def _nhs_check(number: str) -> bool:
    """NHS number modulus-11 check digit validation."""
    digits = [int(d) for d in number if d.isdigit()]
    if len(digits) != 10:
        return False
    total = sum(d * (10 - i) for i, d in enumerate(digits[:9]))
    remainder = total % 11
    check = 11 - remainder
    if check == 11:
        check = 0
    if check == 10:
        return False  # mathematically invalid NHS number
    return check == digits[9]


# ─── Structured PII patterns ──────────────────────────────────────────────────

# Each entry: (rule_id, category, severity, compiled_regex, optional validator)
_PATTERNS: list[tuple[str, str, str, re.Pattern, Optional[Callable[[str], bool]]]] = []

# Static confidence values per rule_id (0.0–1.0).
# Higher = pattern is more specific / has a checksum / has strong context.
# Used in _scan_content(); Presidio uses the live r.score instead.
_RULE_CONFIDENCE: dict[str, float] = {
    "private_key_pem":          0.99,
    "pii_credit_card":          0.95,
    "pii_nhs_number":           0.95,
    "db_conn_string":           0.92,
    "pii_ssn":                  0.90,
    "pii_ni_number":            0.90,
    "jwt_token":                0.90,
    "hardcoded_password":       0.88,
    "pii_email":                0.88,
    "pii_iban":                 0.85,
    "pii_uk_driving_licence":   0.85,
    "pii_uk_sort_code":         0.85,
    "pii_uk_account_number":    0.85,
    "pii_uk_postcode":          0.82,
    "pii_person_name":          0.80,
    "pii_phone":                0.80,
    "pii_uk_phone_mobile":      0.80,
    "pii_phone_intl":           0.75,
    "pii_dob":                  0.78,
    "pii_dob_uk":               0.78,
    "pii_passport":             0.72,
    "pii_mac_address":          0.70,
    "pii_ip_address":           0.65,
    "pii_person_name_bare_key": 0.60,
    "pii_ipv6_address":         0.60,
}


def _p(rule_id: str, category: str, severity: str, pattern: str) -> None:
    _PATTERNS.append((rule_id, category, severity, re.compile(pattern), None))


def _p_luhn(rule_id: str, category: str, severity: str, pattern: str) -> None:
    _PATTERNS.append((rule_id, category, severity, re.compile(pattern), _luhn))


def _p_nhs(rule_id: str, category: str, severity: str, pattern: str) -> None:
    _PATTERNS.append((rule_id, category, severity, re.compile(pattern), _nhs_check))


# Email addresses
_p("pii_email", "pii_email", "medium",
   r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")

# US phone numbers (+1 optional, various separators)
_p("pii_phone", "pii_phone", "medium",
   r"\b(\+1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b")

# International E.164 format (+CC followed by digits)
_p("pii_phone_intl", "pii_phone", "medium",
   r"\+(?!1\b)[1-9]\d{6,14}\b")

# US Social Security Number — excludes invalid prefixes
_p("pii_ssn", "pii_ssn", "critical",
   r"\b(?!000|666|9\d{2})\d{3}[-\s](?!00)\d{2}[-\s](?!0000)\d{4}\b")

# Credit card numbers — Visa, Mastercard, Amex, Discover, Diners (Luhn-validated)
_p_luhn("pii_credit_card", "pii_credit_card", "critical",
        r"\b(?:4\d{12}(?:\d{3})?|5[1-5]\d{14}|3[47]\d{13}|3(?:0[0-5]|[68]\d)\d{11}|6(?:011|5\d{2})\d{12})\b")

# IBAN (basic format check — 2 letters, 2 digits, 1-30 alphanumeric)
_p("pii_iban", "pii_iban_bank", "high",
   r"\b[A-Z]{2}\d{2}[A-Z0-9]{4,30}\b")

# UK National Insurance Number
_p("pii_ni_number", "pii_ssn", "high",
   r"\b[A-CEGHJ-PR-TW-Z]{1}[A-CEGHJ-NPR-TW-Z]{1}\d{6}[A-D]\b")

# Passport number (generic 7–9 alphanumeric)
_p("pii_passport", "pii_passport", "high",
   r"\b[A-Z]{1,2}\d{6,9}\b")

# Date of birth patterns (YYYY-MM-DD, DD/MM/YYYY, MM/DD/YYYY)
# Negative lookbehind/ahead for '/' prevents matching URL archive paths like /2007/07/27/
_p("pii_dob", "pii_date_of_birth", "medium",
   r"(?<!/)(?:19|20)\d{2}[-/](0[1-9]|1[0-2])[-/](0[1-9]|[12]\d|3[01])\b(?!/)")

# PEM private key headers
_p("private_key_pem", "private_key_rsa", "critical",
   r"-----BEGIN (?:RSA |EC |DSA |OPENSSH )?PRIVATE KEY-----")

# Database connection strings with embedded credentials
_p("db_conn_string", "db_connection_string", "critical",
   r"(?:mysql|postgresql|postgres|mongodb|mssql|redis|mariadb)://[^:'\"\s]+:[^@'\"\s]+@[^\s'\"]+")

# Generic password in assignment (name = "value" style)
_p("hardcoded_password", "hardcoded_password", "critical",
   r"""(?i)(?:password|passwd|pwd|secret|pass)\s*[=:]\s*['"][^'"]{4,}['"]""")

# JWT tokens (header.payload.signature)
_p("jwt_token", "jwt_token", "high",
   r"\beyJ[A-Za-z0-9+/\-_]+\.eyJ[A-Za-z0-9+/\-_]+\.[A-Za-z0-9+/\-_]+\b")

# IPv4 addresses inside string literals (potential GDPR personal data)
_p("pii_ip_address", "pii_ip_address", "low",
   r"""['"][0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3}['"]""")

# Full person name in a string literal assigned to a name-related identifier.
# Catches patterns like:  var custName = "Ryan Corcoran";
#                         "customerName": "Jane Smith",
#                         ContactName = 'Alice Brown'
# Two or more Title-case words; key/variable must hint at a person.
_p("pii_person_name", "pii_person_name", "medium",
   r"""(?i)(?:full[_\s]?name|first[_\s]?name|last[_\s]?name|sur[_\s]?name|"""
   r"""given[_\s]?name|display[_\s]?name|contact[_\s]?name|"""
   r"""cust(?:omer)?[_\s]?name|client[_\s]?name|user[_\s]?name|"""
   r"""person[_\s]?name|owner[_\s]?name|staff[_\s]?name|employee[_\s]?name|"""
   r"""author[_\s]?name|applicant[_\s]?name|recipient[_\s]?name)"""
   r"""\s*[=:]\s*['"][A-Z][a-z]{1,30}(?:\s[A-Z][a-z]{1,30}){1,3}['"]""")

# Bare "name" key — JSON {"name": "John Smith"}, YAML `name: John Smith`,
# SQL/Python `name = "John Smith"`.  \bname\b ensures it is a standalone word
# (not inside filename, classname, etc. — underscore counts as \w so var_name
# has no word-boundary before the 'n').  ['"]? before the separator handles
# the closing quote of a JSON key ("name":).
_p("pii_person_name_bare_key", "pii_person_name", "medium",
   r"""(?i)\bname\b['"]?\s*[=:]\s*['"]?[A-Z][a-z]{1,30}(?:\s[A-Z][a-z]{1,30}){1,3}['"]?""")

# ─── UK-specific PII patterns ─────────────────────────────────────────────────

# UK postcode — covers all Royal Mail formats (AN, ANN, AAN, AANN, ANA, AANA)
_p("pii_uk_postcode", "pii_uk_postcode", "medium",
   r"\b(?:[A-Z]{1,2}\d{1,2}|[A-Z]{1,2}\d[A-Z])\s*\d[A-Z]{2}\b")

# UK mobile (07xxx xxxxxx) — more specific than the generic E.164 rule
_p("pii_uk_phone_mobile", "pii_phone", "medium",
   r"\b(?:0|\+44\s?|0044\s?)7\d{3}[\s\-]?\d{6}\b")

# DVLA driving licence — 5 surname chars (9-padded) + 6 DOB digits
# + 2 arbitrary chars + 1 digit + 2 letters = 16 chars
_p("pii_uk_driving_licence", "pii_drivers_license", "high",
   r"\b[A-Z9]{5}\d{6}[A-Z0-9]{2}\d[A-Z]{2}\b")

# NHS number — 3-3-4 digit format, modulus-11 check digit validated
_p_nhs("pii_nhs_number", "pii_nhs_number", "high",
       r"\b\d{3}[- ]\d{3}[- ]\d{4}\b")

# UK bank sort code — keyword context required to avoid false positives
# on version strings and other dashed number triples
_p("pii_uk_sort_code", "pii_bank_account", "high",
   r"""(?i)(?:sort[-_\s]?code|sortcode)\s*[=:'"]{0,2}\s*\d{2}[-\s]\d{2}[-\s]\d{2}""")

# UK bank account number (8 digits) — keyword context required
_p("pii_uk_account_number", "pii_bank_account", "high",
   r"""(?i)(?:account[-_\s]?(?:number|no\.?|num)|acct[-_\s]?(?:no\.?|num|number))\s*[=:'"]{0,2}\s*['"]?\d{8}['"]?""")

# Date of birth — DD/MM/YYYY (UK format, complements the YYYY-MM-DD rule)
_p("pii_dob_uk", "pii_date_of_birth", "medium",
   r"\b(0[1-9]|[12]\d|3[01])/(0[1-9]|1[0-2])/(19|20)\d{2}\b")

# IPv6 address — full 8-group colon-hex form
_p("pii_ipv6_address", "pii_ip_address", "low",
   r"\b[0-9a-fA-F]{1,4}(?::[0-9a-fA-F]{1,4}){7}\b")

# MAC address — colon or hyphen separated
_p("pii_mac_address", "pii_mac_address", "low",
   r"\b(?:[0-9A-Fa-f]{2}[:\-]){5}[0-9A-Fa-f]{2}\b")

# ─── File filtering ───────────────────────────────────────────────────────────

_SKIP_DIRS = {
    ".git", "node_modules", "__pycache__", ".venv", "venv", ".tox",
    "dist", "build", ".mypy_cache", ".pytest_cache", "coverage",
    "vendor", "third_party",
}

_TEXT_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".cs", ".cpp", ".c",
    ".h", ".go", ".rb", ".php", ".swift", ".kt", ".rs", ".scala",
    ".sh", ".bash", ".zsh", ".ps1", ".psm1", ".yml", ".yaml", ".json",
    ".xml", ".toml", ".ini", ".cfg", ".conf", ".env", ".properties",
    ".tf", ".tfvars", ".hcl", ".sql", ".md", ".txt", ".html", ".htm",
    ".vue", ".svelte", ".ejs", ".jsx",
    ".csv", ".tsv", ".psv",
}

# Source code file extensions — NER runs only on comment/string-literal chunks
# (already the case via _text_chunks_from_line) but at a higher confidence
# threshold (0.80) to suppress false positives from code identifiers and
# variable names that superficially resemble person names.
_CODE_EXTENSIONS: set[str] = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".cs", ".cpp", ".c",
    ".h", ".go", ".rb", ".php", ".swift", ".kt", ".rs", ".scala",
    ".sh", ".bash", ".zsh", ".ps1", ".psm1",
    ".tf", ".tfvars", ".hcl", ".sql",
    ".vue", ".svelte", ".ejs",
}

# NER confidence thresholds by file class
_NER_THRESHOLD_CODE: float = 0.80   # higher bar for source code files
_NER_THRESHOLD_PROSE: float = 0.70  # standard threshold for prose/data files

_MAX_FILE_SIZE = 2 * 1024 * 1024  # 2 MB — skip very large files


def _should_scan(path: Path) -> bool:
    if any(part in _SKIP_DIRS for part in path.parts):
        return False
    if path.suffix.lower() not in _TEXT_EXTENSIONS and path.suffix != "":
        return False
    try:
        if path.stat().st_size > _MAX_FILE_SIZE:
            return False
    except OSError:
        return False
    return True


# ─── spaCy NER (optional) ─────────────────────────────────────────────────────

_NLP = None
_SPACY_AVAILABLE = False

try:
    import spacy  # type: ignore

    try:
        _NLP = spacy.load("en_core_web_sm", disable=["parser", "tagger", "attribute_ruler", "lemmatizer"])
        _SPACY_AVAILABLE = True
    except OSError:
        print("[pii_scanner] spaCy model 'en_core_web_sm' not found. "
              "Run: python -m spacy download en_core_web_sm", file=sys.stderr)
except ImportError:
    pass  # spaCy is optional

# ─── Presidio NER (optional — preferred over spaCy when both present) ─────────

_PRESIDIO_ANALYZER = None
_PRESIDIO_AVAILABLE = False

try:
    from presidio_analyzer import AnalyzerEngine  # type: ignore
    _PRESIDIO_ANALYZER = AnalyzerEngine()
    _PRESIDIO_AVAILABLE = True
except ImportError:
    pass  # Presidio is optional — spaCy used as fallback when available

# Extracts content from inside string literals only (single or double quoted)
_STRING_LITERAL_RE = re.compile(r'"([^"]{6,}?)"|\'([^\'"]{6,}?)\'')
# Comment text (everything after the comment marker)
_COMMENT_RE = re.compile(r'(?://|#|<!--)\s*(.{6,})')

# Matches the start of a single-line comment for stripping purposes
_COMMENT_STRIP_RE = re.compile(r'(?://|#|<!--)')
# Matches inline block comments /* ... */ on a single line
_INLINE_BLOCK_COMMENT_RE = re.compile(r'/\*.*?\*/')

# A real name: 2–4 words, each Title-case letters only, no digits/underscores
_NAME_RE = re.compile(r'^[A-Z][a-z]{1,30}(?:\s[A-Z][a-z]{1,30}){1,3}$')

# Same but allows a single-word name — used when a column header provides
# strong context (e.g. a CSV "Name" column), reducing false positives.
_NAME_RE_LABELED = re.compile(r'^[A-Z][a-z]{1,30}(?:\s[A-Z][a-z]{1,30}){0,3}$')

# Column headers that reliably indicate a person-name column.
# Matches bare "name" as well as all the compound variants.
_NAME_COLUMN_RE = re.compile(
    r"""^(?:(?:full[_\s]?)?name|first[_\s]?name|last[_\s]?name|sur[_\s]?name|"""
    r"""given[_\s]?name|display[_\s]?name|contact[_\s]?name|"""
    r"""cust(?:omer)?[_\s]?name|client[_\s]?name|person[_\s]?name|"""
    r"""owner[_\s]?name|staff[_\s]?name|employee[_\s]?name|"""
    r"""author[_\s]?name|applicant[_\s]?name|recipient[_\s]?name)$""",
    re.IGNORECASE,
)

# Broader headers that COULD contain person names (e.g. PXCREATEOPNAME,
# AssignedTo, UpdatedBy).  Values are only flagged after NER confirmation
# to keep the false-positive rate low.
_BROAD_NAME_COLUMN_RE = re.compile(
    r"""(?i)(operator|author|assign(?:ee|edto)?|owner|manager|reviewer|"""
    r"""creat(?:e|ed|or|opname|edby)|updat(?:e|ed|or|opname|edby)|"""
    r"""contact|client|customer|staff|employee|"""
    r"""recipient|approver|requestor|handler|submitt(?:er|ed|edby)?|"""
    r"""modifi(?:er|ed|edby)?)""",
)

# Columns that must never be treated as name columns regardless of their label
# (system IDs, checksums, stream blobs, namespace strings).
_SKIP_NAME_COL_RE = re.compile(
    r"""(?i)(systemid|hostid|insid|inskey|pzinskey|checksum|"""
    r"""pzpvstream|streamname|namespace|rulesetname|objclass|pzrule|classname)""",
)


def _ner_confirm_person(text: str, threshold: float = 0.65) -> bool:
    """Return True if NER identifies *text* as a person name.

    Tries Presidio first (score-based), falls back to spaCy when Presidio is
    not installed.  Returns False if neither backend is available.
    """
    text = text.strip()
    if not text:
        return False
    if _PRESIDIO_AVAILABLE and _PRESIDIO_ANALYZER is not None:
        results = _PRESIDIO_ANALYZER.analyze(
            text=text[:500], language="en", entities=["PERSON"]
        )
        return any(r.score >= threshold for r in results)
    if _SPACY_AVAILABLE and _NLP is not None:
        doc = _NLP(text[:300])
        return any(
            ent.label_ == "PERSON" and _NAME_RE.match(ent.text)
            for ent in doc.ents
        )
    return False


def _strip_comments_from_content(content: str) -> str:
    """Return a copy of *content* with comment portions replaced by spaces.

    Handles single-line comment markers (// # <!--) and inline block comments
    (/* ... */).  Multi-line block comments are not stripped.  Character offsets
    are preserved (replaced with spaces) so line numbers remain accurate.
    """
    lines = content.split('\n')
    result = []
    for line in lines:
        # Strip inline block comments first (/* ... */ on the same line)
        line = _INLINE_BLOCK_COMMENT_RE.sub(lambda m: ' ' * len(m.group(0)), line)
        # Then strip from the first single-line comment marker to end of line
        m = _COMMENT_STRIP_RE.search(line)
        if m:
            line = line[:m.start()] + ' ' * (len(line) - m.start())
        result.append(line)
    return '\n'.join(result)


def _text_chunks_from_line(line: str, skip_comments: bool = False) -> list[str]:
    """Return the extracted string literal contents and comment text from a line."""
    chunks: list[str] = []
    for m in _STRING_LITERAL_RE.finditer(line):
        text = (m.group(1) or m.group(2) or "").strip()
        if text:
            chunks.append(text)
    if not skip_comments:
        cm = _COMMENT_RE.search(line)
        if cm:
            chunks.append(cm.group(1).strip())
    return chunks


def _extract_ner_findings(file_rel: str, content: str, show_secrets: bool = False, skip_comments: bool = False) -> list[Finding]:
    """Run NER on string literal / comment text; flag plausible PERSON entities.

    Uses Presidio when available (score-based), falls back to spaCy (low
    severity).  CSV/TSV/PSV files are skipped here because _scan_csv_columns
    handles them with column-context awareness.

    Confidence threshold is raised to _NER_THRESHOLD_CODE (0.80) for source
    code files (see _CODE_EXTENSIONS) to suppress false positives from code
    identifiers that resemble names.  Comment and string-literal text from
    code files is still fully scanned — author tags, TODO comments, and
    hardcoded names in string literals are all covered.
    """
    suffix = Path(file_rel).suffix.lower()
    if suffix in _CSV_EXTENSIONS:
        return []
    if not _PRESIDIO_AVAILABLE and (not _SPACY_AVAILABLE or _NLP is None):
        return []

    threshold = _NER_THRESHOLD_CODE if suffix in _CODE_EXTENSIONS else _NER_THRESHOLD_PROSE

    # Collect all chunks first so we can skip NER on files with very little
    # extractable prose — avoids paying model inference cost on files that are
    # almost entirely code tokens (imports, braces, short identifiers).
    line_chunks: list[tuple[int, str]] = []  # (lineno, chunk)
    for lineno, line in enumerate(content.splitlines(), start=1):
        for chunk in _text_chunks_from_line(line, skip_comments=skip_comments):
            line_chunks.append((lineno, chunk))

    # Minimum 3 chunks AND at least one chunk with 2+ words before paying NER cost.
    has_prose = any(
        len(chunk.split()) >= 2
        for _, chunk in line_chunks
    )
    if len(line_chunks) < 3 or not has_prose:
        return []

    findings: list[Finding] = []
    for lineno, chunk in line_chunks:
        if _PRESIDIO_AVAILABLE and _PRESIDIO_ANALYZER is not None:
            results = _PRESIDIO_ANALYZER.analyze(
                text=chunk[:500], language="en", entities=["PERSON"]
            )
            for r in results:
                if r.score < threshold:
                    continue
                name = chunk[r.start:r.end].strip()
                if not _NAME_RE.match(name):
                    continue
                category = "pii_person_name"
                rule = _ENGINE.lookup(category)
                findings.append(Finding(
                    id=Finding.make_id(file_rel, lineno, "presidio_person"),
                    scanners=["pii"],
                    category=category,
                    severity="medium",
                    confidence=round(r.score, 4),
                    file=file_rel,
                    line=lineno,
                    match=name if show_secrets else Finding.redact(name),
                    rule_id="presidio_person_name",
                    message=f"Person name detected (Presidio, score={r.score:.2f}): "
                            f"{name if show_secrets else Finding.redact(name)}",
                    remediation_description=rule.description if rule else "",
                    fix_steps=rule.fix_steps if rule else [],
                    references=rule.references if rule else [],
                    regulations=_REG_ENGINE.lookup(category),
                ))
        elif _SPACY_AVAILABLE and _NLP is not None:
            doc = _NLP(chunk[:300])
            for ent in doc.ents:
                if ent.label_ != "PERSON":
                    continue
                if not _NAME_RE.match(ent.text):
                    continue
                category = "pii_person_name"
                rule = _ENGINE.lookup(category)
                findings.append(Finding(
                    id=Finding.make_id(file_rel, lineno, "spacy_person"),
                    scanners=["pii"],
                    category=category,
                    severity="low",
                    confidence=0.65,
                    file=file_rel,
                    line=lineno,
                    match=ent.text if show_secrets else Finding.redact(ent.text),
                    rule_id="spacy_person_name",
                    message=f"Possible person name detected: "
                            f"{ent.text if show_secrets else Finding.redact(ent.text)}",
                    remediation_description=rule.description if rule else "",
                    fix_steps=rule.fix_steps if rule else [],
                    references=rule.references if rule else [],
                    regulations=_REG_ENGINE.lookup(category),
                ))

    return findings


# ─── CSV / TSV column-context name scanner ────────────────────────────────────

_CSV_EXTENSIONS = {".csv", ".tsv", ".psv"}


def _scan_csv_columns(file_rel: str, content: str, show_secrets: bool) -> list[Finding]:
    """Scan CSV/TSV/PSV files using column-header context to locate person names.

    Two-tier matching:
    - *Strict columns* (headers matching _NAME_COLUMN_RE, e.g. "Name", "Full Name"):
      every Title-case value is flagged — medium severity, no NER required.
    - *Broad columns* (headers matching _BROAD_NAME_COLUMN_RE, e.g. PXCREATEOPNAME,
      AssignedTo, CreatedBy): each cell is confirmed with NER (Presidio or spaCy)
      before flagging — medium severity.  Requires Presidio or spaCy to be installed.

    Handles files that begin with blank lines or separator rows (common in Pega /
    Oracle exports) by locating the real header row before parsing.
    """
    lines_raw = content.splitlines()
    if not lines_raw:
        return []

    ext = Path(file_rel).suffix.lower()

    # ── Delimiter detection ────────────────────────────────────────────────────
    # Use the first non-blank, non-separator line so that exports whose first
    # line is blank (e.g. Pega pipe-delimited CSVs) are handled correctly.
    first_meaningful = next(
        (l for l in lines_raw
         if l.strip() and not set(l.strip()) <= {"-", "|", "=", " "}),
        "",
    )
    if ext == ".tsv":
        delimiter = "\t"
    elif ext == ".psv":
        delimiter = "|"
    else:
        candidates = {
            ",": first_meaningful.count(","),
            "\t": first_meaningful.count("\t"),
            "|": first_meaningful.count("|"),
        }
        delimiter = max(candidates, key=lambda d: candidates[d])
        if candidates[delimiter] == 0:
            delimiter = ","

    # ── Parse ──────────────────────────────────────────────────────────────────
    try:
        rows = list(csv.reader(io.StringIO(content), delimiter=delimiter))
    except Exception:
        rows = []

    # ── Find the real header row ───────────────────────────────────────────────
    # Skip blank rows and separator rows (cells that are entirely dashes/pipes).
    header_row_idx = 0
    for i, row in enumerate(rows):
        if row and any(
            c.strip() and not set(c.strip()) <= {"-", "|", "="}
            for c in row
        ):
            header_row_idx = i
            break

    if header_row_idx >= len(rows) - 1:
        return []  # no data rows after header

    headers = rows[header_row_idx]

    # Fallback: if csv.reader badly misaligns columns (>25% of data rows have
    # wrong column count), re-parse with a plain split.
    data_sample = [
        r for r in rows[header_row_idx + 1:]
        if r and not all(set(c.strip()) <= {"-", "|", "="} for c in r)
    ]
    if data_sample:
        expected_cols = len(headers)
        misaligned = sum(1 for r in data_sample if len(r) != expected_cols)
        if misaligned > max(1, len(data_sample) // 4):
            rows = [line.split(delimiter) for line in lines_raw if line.strip()]
            header_row_idx = 0
            for i, row in enumerate(rows):
                if row and any(
                    c.strip() and not set(c.strip()) <= {"-", "|", "="}
                    for c in row
                ):
                    header_row_idx = i
                    break
            headers = rows[header_row_idx]

    if len(rows) <= header_row_idx + 1:
        return []

    # ── Identify columns ───────────────────────────────────────────────────────
    strict_col_indices: list[int] = []
    broad_col_indices: list[int] = []
    for i, h in enumerate(headers):
        h_s = h.strip()
        if _SKIP_NAME_COL_RE.search(h_s):
            continue
        if _NAME_COLUMN_RE.match(h_s):
            strict_col_indices.append(i)
        elif _BROAD_NAME_COLUMN_RE.search(h_s):
            broad_col_indices.append(i)

    ner_available = _PRESIDIO_AVAILABLE or (_SPACY_AVAILABLE and _NLP is not None)
    if not strict_col_indices and (not broad_col_indices or not ner_available):
        return []

    findings: list[Finding] = []

    for row_idx, row in enumerate(rows[header_row_idx + 1:], start=header_row_idx + 2):
        # Skip separator rows in the data
        if not row or all(set(c.strip()) <= {"-", "|", "="} for c in row):
            continue

        # ── Strict columns ─────────────────────────────────────────────────────
        for col_idx in strict_col_indices:
            if col_idx >= len(row):
                continue
            value = row[col_idx].strip()
            if not _NAME_RE_LABELED.match(value):
                continue
            col_header = headers[col_idx].strip()
            category = "pii_person_name"
            rule = _ENGINE.lookup(category)
            findings.append(Finding(
                id=Finding.make_id(file_rel, row_idx, "pii_person_name_column"),
                scanners=["pii"],
                category=category,
                severity="medium",
                file=file_rel,
                line=row_idx,
                match=value if show_secrets else Finding.redact(value),
                rule_id="pii_person_name_column",
                message=f"Person name in '{col_header}' column.",
                remediation_description=rule.description if rule else "",
                fix_steps=rule.fix_steps if rule else [],
                references=rule.references if rule else [],
                regulations=_REG_ENGINE.lookup(category),
            ))

        # ── Broad columns (NER-confirmed) ──────────────────────────────────────
        if not ner_available:
            continue
        for col_idx in broad_col_indices:
            if col_idx >= len(row):
                continue
            value = row[col_idx].strip()
            if len(value) < 4:
                continue
            # Cheap structural pre-filter: must have at least two capitalised
            # words before paying the cost of an NER call.
            if not re.search(r'[A-Z][a-z]{1,25}\s[A-Z][a-z]{1,25}', value):
                continue
            if not _ner_confirm_person(value):
                continue
            col_header = headers[col_idx].strip()
            category = "pii_person_name"
            rule = _ENGINE.lookup(category)
            findings.append(Finding(
                id=Finding.make_id(file_rel, row_idx, f"ner_broad_{col_idx}"),
                scanners=["pii"],
                category=category,
                severity="medium",
                file=file_rel,
                line=row_idx,
                match=value if show_secrets else Finding.redact(value),
                rule_id="pii_person_name_operator_col",
                message=f"Person name in '{col_header}' column (NER-confirmed).",
                remediation_description=rule.description if rule else "",
                fix_steps=rule.fix_steps if rule else [],
                references=rule.references if rule else [],
                regulations=_REG_ENGINE.lookup(category),
            ))

    return findings


# ─── Binary document text extractors ─────────────────────────────────────────

def _extract_docx_text(data: bytes) -> Optional[str]:
    """Extract plain text from a .docx file (requires python-docx)."""
    if not _DOCX_AVAILABLE:
        return None
    try:
        doc = _docx_module.Document(io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as exc:
        print(f"[pii_scanner] docx extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_xlsx_text(data: bytes) -> Optional[str]:
    """Extract cell values from all sheets of an .xlsx file (requires openpyxl)."""
    if not _XLSX_AVAILABLE:
        return None
    try:
        wb = _openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_cells = [str(c) for c in row if c is not None and str(c).strip()]
                if row_cells:
                    parts.append("\t".join(row_cells))
        return "\n".join(parts)
    except Exception as exc:
        print(f"[pii_scanner] xlsx extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_rtf_text(data: bytes) -> Optional[str]:
    """Extract plain text from an .rtf file (requires striprtf)."""
    if not _RTF_AVAILABLE:
        return None
    try:
        return _rtf_to_text(data.decode("utf-8", errors="replace"))
    except Exception as exc:
        print(f"[pii_scanner] rtf extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_eml_text(data: bytes) -> Optional[str]:
    """Extract headers and body text from an .eml email file (stdlib only)."""
    import email as _email_mod
    import email.policy
    try:
        msg = _email_mod.message_from_bytes(data, policy=email.policy.default)
        parts: list[str] = []
        for header in ("From", "To", "Cc", "Bcc", "Reply-To", "Subject", "X-Sender"):
            val = msg.get(header, "")
            if val:
                parts.append(f"{header}: {val}")
        for part in msg.walk():
            if part.get_content_type() in ("text/plain", "text/html"):
                try:
                    payload = part.get_payload(decode=True)
                    if payload:
                        charset = part.get_content_charset() or "utf-8"
                        parts.append(payload.decode(charset, errors="replace"))
                except Exception:
                    pass
        return "\n".join(parts)
    except Exception as exc:
        print(f"[pii_scanner] eml extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_xls_text(data: bytes) -> Optional[str]:
    """Extract cell values from all sheets of a legacy .xls file (requires xlrd)."""
    if not _XLS_AVAILABLE:
        return None
    try:
        wb = _xlrd.open_workbook(file_contents=data)
        parts: list[str] = []
        for sheet in wb.sheets():
            for row_idx in range(sheet.nrows):
                row_cells = [
                    str(sheet.cell_value(row_idx, col))
                    for col in range(sheet.ncols)
                    if str(sheet.cell_value(row_idx, col)).strip()
                ]
                if row_cells:
                    parts.append("\t".join(row_cells))
        return "\n".join(parts)
    except Exception as exc:
        print(f"[pii_scanner] xls extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_doc_text(data: bytes) -> Optional[str]:
    """Extract text from a legacy Word .doc file (requires olefile).

    Reads every OLE stream and extracts printable UTF-16 LE and ASCII runs,
    which is sufficient to surface PII patterns without a full parser.
    """
    if not _DOC_AVAILABLE:
        return None
    try:
        parts: list[str] = []
        with _olefile.OleFileIO(io.BytesIO(data)) as ole:
            for entry in ole.listdir(streams=True):
                try:
                    raw = ole.openstream(entry).read()
                except Exception:
                    continue
                # UTF-16 LE (most Word 97-2003 text streams)
                try:
                    decoded = raw.decode("utf-16-le", errors="replace")
                    parts.extend(re.findall(r'[\x20-\x7e\u00a0-\u024f]{4,}', decoded))
                except Exception:
                    pass
                # ASCII runs as a secondary pass
                parts.extend(
                    s.decode("ascii")
                    for s in re.findall(rb'[\x20-\x7e]{5,}', raw)
                )
        return "\n".join(parts) if parts else None
    except Exception as exc:
        print(f"[pii_scanner] doc extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_pdf_text(data: bytes) -> Optional[str]:
    """Extract the text layer from a text-based PDF (requires pdfminer.six)."""
    if not _PDF_AVAILABLE:
        return None
    try:
        return _pdfminer_extract(io.BytesIO(data))
    except Exception as exc:
        print(f"[pii_scanner] pdf extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_msg_text(data: bytes) -> Optional[str]:
    """Extract headers and body from an Outlook .msg file (requires extract-msg)."""
    if not _MSG_AVAILABLE:
        return None
    try:
        msg = _extract_msg_lib.openMsg(io.BytesIO(data))
        parts: list[str] = []
        for attr in ("sender", "to", "cc", "bcc", "subject"):
            val = getattr(msg, attr, None) or ""
            if val:
                parts.append(f"{attr.capitalize()}: {val}")
        body = msg.body or ""
        if body:
            parts.append(body)
        return "\n".join(parts)
    except Exception as exc:
        print(f"[pii_scanner] msg extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_parquet_text(data: bytes) -> Optional[str]:
    """Extract all cell values from a Parquet file as text (requires pyarrow)."""
    if not _PYARROW_AVAILABLE:
        return None
    try:
        table = _pq.read_table(io.BytesIO(data))
        parts: list[str] = []
        for batch in table.to_batches(max_chunksize=500):
            for col in batch.columns:
                parts.extend(s for v in col.to_pylist() if v is not None and (s := str(v).strip()))
        return "\n".join(parts)
    except Exception as exc:
        print(f"[pii_scanner] parquet extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_orc_text(data: bytes) -> Optional[str]:
    """Extract all cell values from an ORC file as text (requires pyarrow)."""
    if not _PYARROW_AVAILABLE:
        return None
    try:
        table = _pa_orc.ORCFile(io.BytesIO(data)).read()
        parts: list[str] = []
        for col in table.columns:
            parts.extend(s for v in col.to_pylist() if v is not None and (s := str(v).strip()))
        return "\n".join(parts)
    except Exception as exc:
        print(f"[pii_scanner] orc extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_avro_text(data: bytes) -> Optional[str]:
    """Deserialise an Avro container and extract all field values as text (requires fastavro)."""
    if not _AVRO_AVAILABLE:
        return None
    try:
        parts: list[str] = []
        for record in _fastavro.reader(io.BytesIO(data)):
            parts.extend(
                s
                for v in record.values()
                if v is not None and (s := str(v).strip())
            )
        return "\n".join(parts)
    except Exception as exc:
        print(f"[pii_scanner] avro extraction failed: {exc}", file=sys.stderr)
        return None


def _extract_binary_doc(path: Path) -> Optional[str]:
    """Read a binary doc from disk and dispatch to the correct extractor."""
    try:
        data = path.read_bytes()
    except OSError:
        return None
    return _extract_binary_doc_from_bytes(path.suffix.lower(), data)


def _extract_binary_doc_from_bytes(suffix: str, data: bytes) -> Optional[str]:
    """Dispatch bytes to the correct extractor based on file extension."""
    if suffix == ".docx":
        return _extract_docx_text(data)
    if suffix == ".xlsx":
        return _extract_xlsx_text(data)
    if suffix == ".xls":
        return _extract_xls_text(data)
    if suffix == ".rtf":
        return _extract_rtf_text(data)
    if suffix == ".eml":
        return _extract_eml_text(data)
    if suffix == ".doc":
        return _extract_doc_text(data)
    if suffix == ".pdf":
        return _extract_pdf_text(data)
    if suffix == ".msg":
        return _extract_msg_text(data)
    if suffix == ".parquet":
        return _extract_parquet_text(data)
    if suffix == ".orc":
        return _extract_orc_text(data)
    if suffix == ".avro":
        return _extract_avro_text(data)
    return None


# ─── Archive traversal ────────────────────────────────────────────────────────

def _is_archive_path(path: Path) -> bool:
    """Return True if path is a supported archive format."""
    name = path.name.lower()
    return (
        name.endswith((".tar.gz", ".tar.bz2", ".tgz"))
        or path.suffix.lower() in {".zip", ".gz", ".bz2", ".tar"}
    )


def _is_archive_name(name: str) -> bool:
    """Return True if a virtual member name is a supported archive format."""
    lower = name.lower()
    return lower.endswith((".tar.gz", ".tar.bz2", ".tgz", ".zip", ".gz", ".bz2", ".tar"))


def _should_scan_name(name: str) -> bool:
    """Like _should_scan but operates on a virtual name string (archive member)."""
    p = Path(name)
    if any(part in _SKIP_DIRS for part in p.parts):
        return False
    suffix = p.suffix.lower()
    return suffix in _TEXT_EXTENSIONS or suffix in _BINARY_DOC_EXTENSIONS


def _iter_archive_members(archive_name: str, data: bytes) -> Iterator[tuple[str, bytes]]:
    """Yield (member_name, member_bytes) for every file entry in an archive."""
    lower = archive_name.lower()
    buf = io.BytesIO(data)

    try:
        if lower.endswith(".zip"):
            with zipfile.ZipFile(buf) as zf:
                for info in zf.infolist():
                    if not info.is_dir():
                        try:
                            yield info.filename, zf.read(info)
                        except Exception:
                            pass

        elif lower.endswith((".tar.gz", ".tgz")):
            with tarfile.open(fileobj=buf, mode="r:gz") as tf:
                for member in tf.getmembers():
                    if member.isfile():
                        f = tf.extractfile(member)
                        if f:
                            yield member.name, f.read()

        elif lower.endswith(".tar.bz2"):
            with tarfile.open(fileobj=buf, mode="r:bz2") as tf:
                for member in tf.getmembers():
                    if member.isfile():
                        f = tf.extractfile(member)
                        if f:
                            yield member.name, f.read()

        elif lower.endswith(".tar"):
            with tarfile.open(fileobj=buf, mode="r:") as tf:
                for member in tf.getmembers():
                    if member.isfile():
                        f = tf.extractfile(member)
                        if f:
                            yield member.name, f.read()

        elif lower.endswith(".gz"):
            # Single-file gzip (not .tar.gz — handled above)
            inner_name = archive_name[:-3]  # strip .gz suffix
            yield inner_name, gzip.decompress(data)

        elif lower.endswith(".bz2"):
            inner_name = archive_name[:-4]  # strip .bz2 suffix
            yield inner_name, bz2.decompress(data)

    except Exception as exc:
        print(f"[pii_scanner] Could not open archive {archive_name!r}: {exc}", file=sys.stderr)


# ─── Base64 / URL / JWT decode scanning ──────────────────────────────────────

# Matches potential base64 blobs (min 24 chars, standard alphabet + optional padding)
_BASE64_BLOB_RE = re.compile(r'[A-Za-z0-9+/]{24,}={0,2}')
# Matches URL percent-encoded sequences (at least 2 consecutive %XX pairs)
_URL_ENCODED_RE = re.compile(r'(?:%[0-9A-Fa-f]{2}){2,}')
# Captures the JWT payload (second segment, starts with eyJ when decoded)
_JWT_PAYLOAD_RE = re.compile(
    r'\beyJ[A-Za-z0-9+/\-_]+=*\.eyJ([A-Za-z0-9+/\-_]+=*)\.[A-Za-z0-9+/\-_]+=*'
)


def _scan_chunk_for_pii(
    file_rel: str,
    lineno: int,
    text: str,
    show_secrets: bool,
    source_tag: str,
) -> list[Finding]:
    """Run all regex PII patterns against a decoded text chunk."""
    results: list[Finding] = []
    for rule_id, category, severity, pattern, validator in _PATTERNS:
        for match in pattern.finditer(text):
            matched_text = match.group(0)
            if validator and not validator(matched_text):
                continue
            rule = _ENGINE.lookup(category)
            results.append(Finding(
                id=Finding.make_id(file_rel, lineno, f"{rule_id}_{source_tag}"),
                scanners=["pii"],
                category=category,
                severity=severity,
                file=file_rel,
                line=lineno,
                match=matched_text if show_secrets else Finding.redact(matched_text),
                rule_id=rule_id,
                message=(
                    f"{category.replace('_', ' ').title()} detected in "
                    f"{source_tag.replace('_', ' ')} payload."
                ),
                remediation_description=rule.description if rule else "",
                fix_steps=rule.fix_steps if rule else [],
                references=rule.references if rule else [],
                regulations=_REG_ENGINE.lookup(category),
            ))
    return results


def _scan_decoded_payloads(
    file_rel: str, content: str, show_secrets: bool
) -> list[Finding]:
    """Decode Base64, URL-encoded, and JWT payloads found in content, then scan for PII."""
    findings: list[Finding] = []
    for lineno, line in enumerate(content.splitlines(), start=1):

        # ── URL-encoded strings ────────────────────────────────────────────────
        for m in _URL_ENCODED_RE.finditer(line):
            decoded = urllib.parse.unquote(m.group(0))
            if decoded != m.group(0):
                findings.extend(
                    _scan_chunk_for_pii(file_rel, lineno, decoded, show_secrets, "url_decoded")
                )

        # ── JWT payload ────────────────────────────────────────────────────────
        for m in _JWT_PAYLOAD_RE.finditer(line):
            payload_b64 = m.group(1).replace("-", "+").replace("_", "/")
            payload_b64 += "=" * (-len(payload_b64) % 4)
            try:
                decoded = base64.b64decode(payload_b64).decode("utf-8", errors="replace")
                findings.extend(
                    _scan_chunk_for_pii(file_rel, lineno, decoded, show_secrets, "jwt_payload")
                )
            except Exception:
                pass

        # ── Generic Base64 blobs ───────────────────────────────────────────────
        for m in _BASE64_BLOB_RE.finditer(line):
            blob = m.group(0)
            padded = blob + "=" * (-len(blob) % 4)
            try:
                decoded_bytes = base64.b64decode(padded.encode(), validate=True)
                decoded = decoded_bytes.decode("utf-8")  # strict — rejects binary blobs
                if len(decoded) >= 8 and any(c.isprintable() and not c.isspace() for c in decoded[:20]):
                    findings.extend(
                        _scan_chunk_for_pii(file_rel, lineno, decoded, show_secrets, "base64_decoded")
                    )
            except Exception:
                pass  # not valid UTF-8 base64 — skip silently

    return findings


# ─── Scanner class ────────────────────────────────────────────────────────────

class PIIScanner(AbstractScanner):
    @property
    def name(self) -> str:
        return "presidio"

    async def is_available(self) -> bool:
        return True  # Pure Python — always available

    async def scan(self, config: ScanConfig) -> list[Finding]:
        self._files_scanned: int = 0
        self._files_skipped: int = 0
        self._lines_scanned: int = 0
        self._lines_skipped: int = 0
        try:
            return await self._run(config)
        except Exception as exc:
            print(f"[pii_scanner] Scan failed: {exc}", file=sys.stderr)
            return []

    async def _run(self, config: ScanConfig) -> list[Finding]:
        root = Path(config.path).resolve()
        excluded = set(config.exclude_paths)
        excluded_files = set(config.exclude_files)
        exclude_patterns = config.exclude_patterns
        loop = asyncio.get_running_loop()

        # Walk the filesystem in a thread so rglob() does not block the event loop
        all_paths: list[Path] = await loop.run_in_executor(
            None, lambda: list(root.rglob("*"))
        )

        # Process files sequentially — CPU-bound regex and NER do not benefit
        # from concurrency on the event loop thread (they serialise under the GIL
        # and context-switching overhead makes things worse).  Binary-doc extraction
        # and archive reads are still offloaded to the executor inside _process_file.
        findings: list[Finding] = []
        for path in all_paths:
            if any(part in excluded for part in path.parts):
                continue
            try:
                rel = str(path.relative_to(root))
            except ValueError:
                rel = str(path)
            rel_posix = rel.replace("\\", "/")
            if rel in excluded_files or rel_posix in excluded_files:
                continue
            if exclude_patterns and any(fnmatch.fnmatch(rel_posix, pat) for pat in exclude_patterns):
                continue
            if not path.is_file():
                continue
            findings.extend(await self._process_file(path, root, config))
        return findings

    async def _process_file(
        self,
        path: Path,
        root: Path,
        config: ScanConfig,
    ) -> list[Finding]:
        """Scan a single file and return all findings for it.

        Blocking I/O (archive reads, binary document extraction) is offloaded
        to the default thread-pool executor so it does not stall the event loop
        while other files are being processed concurrently.
        """
        loop = asyncio.get_running_loop()
        suffix = path.suffix.lower()

        try:
            file_rel = str(path.relative_to(root))
        except ValueError:
            file_rel = str(path)

        # ── Archives ──────────────────────────────────────────────────────────
        if _is_archive_path(path):
            try:
                data = await loop.run_in_executor(None, path.read_bytes)
                return self._scan_archive_data(file_rel, data, config, depth=0)
            except OSError:
                return []

        # ── Binary documents (.docx, .xlsx, .rtf, .eml, ...) ─────────────────
        if suffix in _BINARY_DOC_EXTENSIONS:
            content = await loop.run_in_executor(None, _extract_binary_doc, path)
            if content is None:
                self._files_skipped += 1
                return []

        else:
            # ── Regular text files ─────────────────────────────────────────────
            if not _should_scan(path):
                self._files_skipped += 1
                return []
            try:
                async with aiofiles.open(path, encoding="utf-8", errors="replace") as fh:
                    content = await fh.read()
            except OSError:
                return []

        findings: list[Finding] = []
        findings.extend(self._scan_content(file_rel, content, config.show_secrets, config.skip_comments))
        findings.extend(_extract_ner_findings(file_rel, content, config.show_secrets, config.skip_comments))
        findings.extend(_scan_decoded_payloads(file_rel, content, config.show_secrets))
        if suffix in _CSV_EXTENSIONS:
            findings.extend(_scan_csv_columns(file_rel, content, config.show_secrets))

        file_lines = content.splitlines()
        self._files_scanned += 1
        self._lines_scanned += len(file_lines)
        if config.skip_comments:
            stripped_lines = _strip_comments_from_content(content).splitlines()
            self._lines_skipped += sum(
                1 for o, s in zip(file_lines, stripped_lines)
                if o.strip() and not s.strip()
            )
        return findings

    def _scan_archive_data(
        self,
        archive_rel: str,
        data: bytes,
        config: ScanConfig,
        depth: int,
    ) -> list[Finding]:
        """Recursively scan the contents of an archive up to _MAX_ARCHIVE_DEPTH."""
        if depth >= _MAX_ARCHIVE_DEPTH:
            print(
                f"[pii_scanner] Skipping {archive_rel!r}: "
                f"max archive depth ({_MAX_ARCHIVE_DEPTH}) exceeded.",
                file=sys.stderr,
            )
            return []

        findings: list[Finding] = []
        for member_name, member_data in _iter_archive_members(archive_rel, data):
            member_rel = f"{archive_rel}/{member_name}"

            # Nested archive — recurse
            if _is_archive_name(member_name):
                findings.extend(
                    self._scan_archive_data(member_rel, member_data, config, depth + 1)
                )
                continue

            # Binary doc — extract then scan
            member_suffix = Path(member_name).suffix.lower()
            if member_suffix in _BINARY_DOC_EXTENSIONS:
                content = _extract_binary_doc_from_bytes(member_suffix, member_data)
                if content is None:
                    continue
            else:
                # Text file — skip unrecognised extensions
                if not _should_scan_name(member_name):
                    continue
                try:
                    content = member_data.decode("utf-8", errors="replace")
                except Exception:
                    continue

            findings.extend(self._scan_content(member_rel, content, config.show_secrets, config.skip_comments))
            findings.extend(_extract_ner_findings(member_rel, content, config.show_secrets, config.skip_comments))
            findings.extend(_scan_decoded_payloads(member_rel, content, config.show_secrets))
            if member_suffix in _CSV_EXTENSIONS:
                findings.extend(_scan_csv_columns(member_rel, content, config.show_secrets))
            member_lines = content.splitlines()
            self._lines_scanned += len(member_lines)
            if config.skip_comments:
                stripped_lines = _strip_comments_from_content(content).splitlines()
                self._lines_skipped += sum(
                    1 for o, s in zip(member_lines, stripped_lines)
                    if o.strip() and not s.strip()
                )
            self._files_scanned += 1

        return findings

    def _scan_content(self, file_rel: str, content: str, show_secrets: bool = False, skip_comments: bool = False) -> list[Finding]:
        findings: list[Finding] = []
        scan_content = _strip_comments_from_content(content) if skip_comments else content

        for rule_id, category, severity, pattern, validator in _PATTERNS:
            for match in pattern.finditer(scan_content):
                matched_text = match.group(0)

                if validator and not validator(matched_text):
                    continue

                # Calculate 1-based line number from character offset
                line = scan_content[: match.start()].count("\n") + 1
                rule = _ENGINE.lookup(category)

                findings.append(Finding(
                    id=Finding.make_id(file_rel, line, rule_id),
                    scanners=["pii"],
                    category=category,
                    severity=severity,
                    confidence=_RULE_CONFIDENCE.get(rule_id, 0.70),
                    file=file_rel,
                    line=line,
                    match=matched_text if show_secrets else Finding.redact(matched_text),
                    rule_id=rule_id,
                    message=f"{category.replace('_', ' ').title()} pattern detected.",
                    remediation_description=rule.description if rule else "",
                    fix_steps=rule.fix_steps if rule else [],
                    references=rule.references if rule else [],
                    regulations=_REG_ENGINE.lookup(category),
                ))

        return findings
